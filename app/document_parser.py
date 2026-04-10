from __future__ import annotations

import re
import unicodedata
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path

import dateparser
import pdfplumber
from docx import Document

from .models import DeliveryItem

DATE_PATTERNS = [
    r"\b\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?\b",
    r"\b\d{1,2}\s+de\s+[^0-9\s]+\s*(?:de\s+\d{4})?\b",
    r"\b\d{1,2}\s+[^0-9\s\.]+\s*(?:\d{4})?\b",
    r"\b[^0-9\s,]+\s+\d{1,2},?\s+\d{4}\b",
]

READING_KEYWORDS = (
    "lectura",
    "lecturas",
    "leer",
    "revision de texto",
    "capitulo",
    "capitulos",
    "articulo",
    "articulos",
    "autor",
    "autores",
    "guia de consulta",
    "doi",
    "revista",
    "disponible en",
)

STRONG_ACTIVITY_KEYWORDS = (
    "actividad evaluativa",
    "actividad integrada",
    "taller",
    "parcial",
    "exposicion",
    "entrega",
    "evaluacion",
    "quiz",
    "examen",
    "analisis de caso",
    "sustentacion",
    "prueba objetiva",
    "informe",
    "papsi",
    "coloquio",
    "pelicula y analisis",
    "proyecto de aula",
    "trabajo final",
    "seminario aleman",
)

WEAK_ACTIVITY_KEYWORDS = (
    "actividad",
    "trabajo",
    "seminario",
)

NOISE_KEYWORDS = (
    "clase magistral",
    "socializacion de consultas",
    "socializacion",
    "discusion sobre el tema",
    "discusion grupal",
    "plantear preguntas orientadoras",
    "establecer elementos integradores",
    "uso de material audiovisual",
    "explicacion general",
    "presentacion del programa",
    "plan de curso",
    "sugerencia de lecturas",
    "temas para desarrollar",
    "semana santa",
    "no hay clase",
)

ACTIVITY_TYPE_PATTERNS = (
    ("actividad evaluativa", "Actividad evaluativa"),
    ("actividad integrada", "Actividad integrada"),
    ("analisis de caso", "Analisis de caso"),
    ("pelicula y analisis", "Pelicula y analisis"),
    ("prueba objetiva", "Prueba objetiva"),
    ("examen final", "Examen final"),
    ("examen", "Examen"),
    ("parcial", "Parcial"),
    ("quiz", "Quiz"),
    ("exposicion", "Exposicion"),
    ("sustentacion", "Sustentacion"),
    ("entrega", "Entrega"),
    ("taller", "Taller"),
    ("informe", "Informe"),
    ("papsi", "PAPsi"),
    ("coloquio", "Coloquio"),
    ("seminario aleman", "Seminario"),
    ("proyecto de aula", "Proyecto de aula"),
    ("trabajo", "Trabajo"),
)

SECTION_LABEL_PATTERN = re.compile(
    r"(?=(?:desarrollo|neuropsicologia|psicopatologia|psicologia clinica)\s*:)",
    flags=re.IGNORECASE,
)

CITATION_PATTERNS = (
    re.compile(
        r"(?:^|[.;])\s*(?P<author>[a-z][a-z\s\-]{2,80}?)(?:,\s*[^()]*)?\(\d{4}\)\.\s*(?P<title>[^.]{6,140})",
        flags=re.IGNORECASE,
    ),
    re.compile(
        r"(?:^|[.;])\s*(?P<author>[a-z][a-z\s,\-&]{2,100}?)\(\d{4}\)\s*[\"'«]?(?P<title>[^.\"'»]{6,140})",
        flags=re.IGNORECASE,
    ),
    re.compile(
        r"(?:^|[.;])\s*(?P<author>asociacion americana de psiquiatria)\s*\(\d{4}\)\.\s*(?P<title>[^.]{6,140})",
        flags=re.IGNORECASE,
    ),
)

GENERIC_DETAIL_PATTERNS = (
    "sobre el tema",
    "articulando los tres componentes",
    "tematicas estudiantes",
    "estudiantes",
    "docente",
    "del curso",
)


def analyze_document(
    filename: str,
    content: bytes,
    today: date | None = None,
    reminder_days: int = 5,
) -> list[DeliveryItem]:
    if today is None:
        today = date.today()

    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return _parse_docx_document(filename, content, today, reminder_days)
    if suffix == ".pdf":
        text = _extract_pdf_text(content)
        return parse_deliveries(text, today=today, reminder_days=reminder_days, source_name=filename)
    raise ValueError("Formato no soportado. Usa un archivo PDF o DOCX.")


def extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(content)
    if suffix == ".docx":
        return _extract_docx_text(content)
    raise ValueError("Formato no soportado. Usa un archivo PDF o DOCX.")


def parse_deliveries(
    text: str,
    today: date | None = None,
    reminder_days: int = 5,
    source_name: str = "",
    base_year: int | None = None,
) -> list[DeliveryItem]:
    if today is None:
        today = date.today()
    reminder_days = max(0, reminder_days)
    subject = _extract_subject(text, source_name)

    deliveries: list[DeliveryItem] = []
    seen: set[tuple[str, str, str]] = set()

    for raw_line in _normalize_lines(text):
        found_date = _extract_date(raw_line, today, base_year=base_year)
        if not found_date:
            continue

        category = _classify_line(raw_line)
        if category == "otro":
            continue

        if category == "lectura":
            title = _summarize_reading_section(raw_line, raw_line)
        else:
            title = _summarize_activity_segment(raw_line, raw_line)

        if not title:
            title = _extract_title(raw_line, found_date["matched_text"], category)
        if not title:
            continue

        item = _build_item(
            subject=subject,
            category=category,
            title=title,
            due_date=found_date["date"],
            source_line=raw_line,
            reminder_days=reminder_days,
        )
        key = (item.subject.lower(), item.category.lower(), f"{item.title.lower()}|{item.due_date_iso}")
        if key not in seen:
            seen.add(key)
            deliveries.append(item)

    deliveries.sort(key=lambda item: (item.due_date_iso, item.subject, item.title))
    return deliveries


def _parse_docx_document(
    filename: str,
    content: bytes,
    today: date,
    reminder_days: int,
) -> list[DeliveryItem]:
    document = Document(BytesIO(content))
    subject = _extract_subject_from_document(document, filename)
    base_year = _extract_year_from_document(document, filename, today.year)
    deliveries = _extract_items_from_tables(document, subject, today, reminder_days, base_year)

    if deliveries:
        return _dedupe_items(deliveries)

    text = _extract_docx_text(content)
    return parse_deliveries(
        text,
        today=today,
        reminder_days=reminder_days,
        source_name=filename,
        base_year=base_year,
    )


def _extract_items_from_tables(
    document: Document,
    subject: str,
    today: date,
    reminder_days: int,
    base_year: int,
) -> list[DeliveryItem]:
    items: list[DeliveryItem] = []
    for table in document.tables:
        rows = [_row_to_cells(row) for row in table.rows]
        if not rows:
            continue

        header = [_normalize_match_text(cell) for cell in rows[0]]
        if any("semana" in cell for cell in header) and any("trabajo independiente" in cell for cell in header):
            items.extend(_parse_course_plan_table(rows, subject, today, reminder_days, base_year))
        elif any("fecha" in cell for cell in header) and any("tipo de prueba" in cell for cell in header):
            items.extend(_parse_evaluation_plan_table(rows, subject, today, reminder_days, base_year))

    return items


def _parse_course_plan_table(
    rows: list[list[str]],
    subject: str,
    today: date,
    reminder_days: int,
    base_year: int,
) -> list[DeliveryItem]:
    items: list[DeliveryItem] = []
    for row in rows[1:]:
        if len(row) < 4:
            continue

        week_cell, themes_cell, teacher_cell, independent_cell = row[:4]
        due_info = _extract_date(week_cell, today, base_year=base_year)
        if not due_info:
            continue

        due_date = due_info["date"]
        row_text = " | ".join(part for part in row if part)
        if "semana santa" in _normalize_match_text(row_text):
            continue

        for title in _extract_activities_from_course_row(themes_cell, teacher_cell, independent_cell):
            items.append(
                _build_item(
                    subject=subject,
                    category="actividad",
                    title=title,
                    due_date=due_date,
                    source_line=row_text,
                    reminder_days=reminder_days,
                )
            )

        for title in _extract_readings_from_course_row(themes_cell, independent_cell):
            items.append(
                _build_item(
                    subject=subject,
                    category="lectura",
                    title=title,
                    due_date=due_date,
                    source_line=row_text,
                    reminder_days=reminder_days,
                )
            )

    return items


def _parse_evaluation_plan_table(
    rows: list[list[str]],
    subject: str,
    today: date,
    reminder_days: int,
    base_year: int,
) -> list[DeliveryItem]:
    items: list[DeliveryItem] = []
    for row in rows[1:]:
        if len(row) < 4:
            continue

        topic, test_type, percent, date_cell = row[:4]
        due_info = _extract_date(date_cell, today, base_year=base_year)
        if not due_info:
            continue

        title = _summarize_evaluation_item(topic, test_type)
        if not title:
            continue

        row_text = " | ".join(part for part in row if part)
        items.append(
            _build_item(
                subject=subject,
                category="actividad",
                title=title,
                due_date=due_info["date"],
                source_line=row_text,
                reminder_days=reminder_days,
            )
        )

    return items


def _extract_activities_from_course_row(themes: str, teacher_work: str, independent_work: str) -> list[str]:
    summaries: list[str] = []
    for text in (teacher_work, independent_work, themes):
        for segment in _split_activity_segments(text):
            summary = _summarize_activity_segment(segment, themes)
            if summary:
                summaries.append(summary)
    return _dedupe_strings(summaries)[:4]


def _extract_readings_from_course_row(themes: str, independent_work: str) -> list[str]:
    normalized_independent = _normalize_match_text(independent_work)
    if not normalized_independent or "no aplica" in normalized_independent or "semana santa" in normalized_independent:
        return []

    sections = _split_course_sections(independent_work)
    candidates: list[str] = []
    for section in sections:
        if not (_contains_any(section, READING_KEYWORDS) or _looks_like_bibliography(section)):
            continue
        summary = _summarize_reading_section(section, themes)
        if summary:
            candidates.append(summary)

    if candidates:
        return _dedupe_strings(candidates)[:3]

    fallback = _extract_theme_hint(themes)
    return [f"Lectura: {fallback}"] if fallback else []


def _summarize_evaluation_item(topic: str, test_type: str) -> str:
    normalized_topic = _normalize_match_text(topic)
    normalized_test_type = _normalize_match_text(test_type)

    topic_type = _detect_activity_type(normalized_topic)
    if topic_type:
        detail = _clean_detail_text(normalized_test_type or normalized_topic.replace(topic_type.lower(), ""))
        if detail and detail != topic_type.lower():
            return _truncate_display(f"{topic_type}: {detail}")
        return topic_type

    test_type_label = _detect_activity_type(normalized_test_type)
    topic_label = _clean_detail_text(normalized_topic)
    if test_type_label and topic_label:
        return _truncate_display(f"{test_type_label}: {topic_label}")

    summary = _summarize_activity_segment(topic, topic)
    return summary or ""


def _summarize_reading_section(section: str, theme_context: str = "") -> str:
    normalized_section = _normalize_match_text(_strip_section_prefix(section))
    if not normalized_section:
        return ""

    candidates = _extract_citation_candidates(normalized_section)
    if candidates:
        author, title = max(
            candidates,
            key=lambda candidate: _score_reading_candidate(candidate[0], candidate[1], theme_context),
        )
        author_display = _display_author(author, title)
        title_display = _clean_reading_title(title)
        if author_display and title_display and author_display.lower() not in title_display.lower():
            return _truncate_display(f"Lectura: {author_display} - {title_display}")
        if title_display:
            return _truncate_display(f"Lectura: {title_display}")

    fallback_title = _extract_first_reading_phrase(normalized_section)
    if fallback_title:
        return _truncate_display(f"Lectura: {fallback_title}")

    if _contains_any(normalized_section, ("lectura", "leer")):
        theme_hint = _extract_theme_hint(theme_context or normalized_section)
        if theme_hint:
            return _truncate_display(f"Lectura: {theme_hint}")

    return ""


def _summarize_activity_segment(segment: str, theme_context: str = "") -> str:
    normalized_segment = _normalize_match_text(_strip_section_prefix(segment))
    if not normalized_segment:
        return ""
    if not _is_actionable_activity_text(normalized_segment):
        return ""
    if _is_noise_only(normalized_segment):
        return ""

    activity_type = _detect_activity_type(normalized_segment)
    if not activity_type:
        return ""

    detail = _extract_activity_detail(normalized_segment, activity_type, theme_context)
    if not detail:
        if activity_type in {
            "Actividad integrada",
            "Proyecto de aula",
            "Trabajo",
            "Examen final",
            "Examen",
            "Prueba objetiva",
            "Parcial",
            "Quiz",
            "Informe",
            "Entrega",
        }:
            return ""
        return activity_type

    if detail.lower() == activity_type.lower():
        return activity_type

    return _truncate_display(f"{activity_type}: {detail}")


def _extract_activity_detail(segment: str, activity_type: str, theme_context: str) -> str:
    detail = ""
    prefix = _normalize_match_text(activity_type)
    matched_keyword = next(
        (keyword for keyword, label in ACTIVITY_TYPE_PATTERNS if label == activity_type and _keyword_in_text(segment, keyword)),
        prefix,
    )

    if ":" in segment:
        before, after = segment.split(":", 1)
        if prefix in before:
            detail = after

    if not detail:
        parenthetical = re.search(r"\(([^()]{4,120})\)", segment)
        if parenthetical:
            detail = parenthetical.group(1)

    if not detail:
        detail = segment
        detail = re.sub(r"\b" + re.escape(matched_keyword).replace(r"\ ", r"\s+") + r"\b", " ", detail, count=1)
        detail = detail.replace(":", " ")

    detail = _clean_detail_text(detail)
    if _is_generic_detail(detail):
        detail = _extract_theme_hint(theme_context)

    if activity_type == "Actividad evaluativa" and not detail and "integrada" in segment:
        detail = "integrada"

    if activity_type == "Actividad integrada" and "visita parque explora" in segment:
        detail = "visita parque explora"

    if activity_type == "Entrega" and "proyecto de aula" in segment:
        detail = "proyecto de aula"

    if activity_type == "Proyecto de aula" and ("asesoria" in detail or not detail):
        return ""

    return _clean_detail_text(detail)


def _split_activity_segments(text: str) -> list[str]:
    segments: list[str] = []
    for section in _split_course_sections(text):
        parts = re.split(r"\s*\|\s*|(?<=[.;])\s+", section)
        for part in parts:
            compact = _clean_detail_text(part)
            if compact:
                segments.append(compact)
    return segments


def _split_course_sections(text: str) -> list[str]:
    normalized = _normalize_match_text(text)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    if not normalized:
        return []
    if not SECTION_LABEL_PATTERN.search(normalized):
        return [normalized]
    return [section.strip(" .;|") for section in SECTION_LABEL_PATTERN.split(normalized) if section.strip(" .;|")]


def _extract_citation_candidates(section: str) -> list[tuple[str, str]]:
    candidates: list[tuple[str, str]] = []
    for pattern in CITATION_PATTERNS:
        for match in pattern.finditer(section):
            author = _clean_detail_text(match.group("author"))
            title = _clean_reading_title(match.group("title"))
            if author and title:
                candidates.append((author, title))
    return candidates


def _score_reading_candidate(author: str, title: str, theme_context: str) -> int:
    score = 0
    normalized_title = _normalize_match_text(title)
    theme_keywords = set(_extract_significant_words(theme_context))

    overlap = sum(1 for word in theme_keywords if word in normalized_title)
    score += min(overlap, 3) * 3

    word_count = len(title.split())
    if 3 <= word_count <= 14:
        score += 2
    if "dsm 5" in normalized_title:
        score += 1
    if author.startswith("asociacion americana") and "guia de consulta" in normalized_title:
        score -= 1
    if normalized_title in {"desarrollo humano", "guia de consulta de los criterios diagnosticos del dsm 5"}:
        score -= 1
    return score


def _display_author(author: str, title: str) -> str:
    normalized_title = _normalize_match_text(title)
    if "dsm 5" in normalized_title:
        return "DSM-5"

    primary = author.split(",")[0].strip()
    primary = re.sub(r"\b(?:y|and)\b.*$", "", primary).strip()
    if not primary:
        return ""
    return " ".join(part.capitalize() for part in primary.split())


def _clean_reading_title(title: str) -> str:
    cleaned = _clean_detail_text(title)
    cleaned = re.sub(r"^\(?\d{4}\)?\s*", "", cleaned)
    cleaned = cleaned.strip("\"'«» ")
    cleaned = re.sub(r"\brecuperado de\b.*$", "", cleaned).strip(" .,:;-")
    cleaned = re.sub(r"\ben [a-z].*$", "", cleaned).strip(" .,:;-")
    return cleaned


def _is_actionable_activity_text(normalized_text: str) -> bool:
    if _contains_any(normalized_text, STRONG_ACTIVITY_KEYWORDS):
        return True
    if "actividad" in normalized_text and any(token in normalized_text for token in ("integrada", "evaluativa", "final")):
        return True
    return False


def _is_noise_only(normalized_text: str) -> bool:
    return _contains_any(normalized_text, NOISE_KEYWORDS) and not _contains_any(
        normalized_text,
        STRONG_ACTIVITY_KEYWORDS,
    )


def _detect_activity_type(normalized_text: str) -> str | None:
    for keyword, label in ACTIVITY_TYPE_PATTERNS:
        if _keyword_in_text(normalized_text, keyword):
            return label
    return None


def _is_generic_detail(detail: str) -> bool:
    if not detail:
        return True
    normalized = _normalize_match_text(detail)
    if re.fullmatch(r"[\d\s%.-]+", normalized):
        return True
    if normalized in {"integrada", "evaluativa", "actividad", "trabajo", "tematicas", "estudiantes"}:
        return True
    return any(pattern in normalized for pattern in GENERIC_DETAIL_PATTERNS)


def _extract_significant_words(text: str) -> list[str]:
    normalized = _normalize_match_text(text)
    words = re.findall(r"[a-z]{4,}", normalized)
    stopwords = {
        "desarrollo",
        "neuropsicologia",
        "psicopatologia",
        "actividad",
        "lectura",
        "adulto",
        "adultez",
        "clase",
        "modulo",
        "tema",
        "temas",
        "integrada",
        "integrado",
        "sugeridas",
        "sobre",
        "para",
    }
    return [word for word in words if word not in stopwords]


def _build_item(
    subject: str,
    category: str,
    title: str,
    due_date: date,
    source_line: str,
    reminder_days: int,
) -> DeliveryItem:
    reminder_date = due_date - timedelta(days=max(0, reminder_days))
    return DeliveryItem(
        subject=subject,
        category=category,
        title=title[:160],
        due_date_iso=due_date.isoformat(),
        reminder_date_iso=reminder_date.isoformat(),
        source_line=source_line[:1000],
        reminder_days=max(0, reminder_days),
    )


def _dedupe_items(items: list[DeliveryItem]) -> list[DeliveryItem]:
    seen: set[tuple[str, str, str, str]] = set()
    unique: list[DeliveryItem] = []
    for item in items:
        key = (
            item.subject.lower(),
            item.category.lower(),
            item.title.lower(),
            item.due_date_iso,
        )
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)
    unique.sort(key=lambda item: (item.due_date_iso, item.subject, item.category, item.title))
    return unique


def _extract_pdf_text(content: bytes) -> str:
    pages: list[str] = []
    with pdfplumber.open(BytesIO(content)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(_row_to_cells(row)))
    return "\n".join(lines)


def _row_to_cells(row) -> list[str]:
    cells: list[str] = []
    for cell in row.cells:
        text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
        normalized = re.sub(r"\s+", " ", text).strip()
        cells.append(normalized)
    return cells


def _normalize_lines(text: str) -> list[str]:
    cleaned = text.replace("\r", "\n")
    chunks = [line.strip(" -•\t") for line in cleaned.split("\n")]
    return [re.sub(r"\s+", " ", chunk).strip() for chunk in chunks if chunk.strip()]


def _extract_date(line: str, today: date, base_year: int | None = None) -> dict | None:
    for pattern in DATE_PATTERNS:
        for match in re.finditer(pattern, line, flags=re.IGNORECASE):
            matched_text = match.group(0).replace(".", "").strip()
            if "semana" in _normalize_match_text(matched_text):
                continue

            parsed = dateparser.parse(
                matched_text,
                languages=["es", "en"],
                settings={
                    "PREFER_DATES_FROM": "future",
                    "RELATIVE_BASE": datetime.combine(today, datetime.min.time()),
                },
            )
            if not parsed:
                continue

            parsed_date = parsed.date()
            explicit_year = re.search(r"\b\d{4}\b", matched_text)
            if base_year and not explicit_year:
                parsed_date = parsed_date.replace(year=base_year)
            elif parsed_date.year < today.year and not explicit_year:
                parsed_date = parsed_date.replace(year=today.year)

            return {"date": parsed_date, "matched_text": matched_text}
    return None


def _classify_line(line: str) -> str:
    normalized = _normalize_match_text(line)
    if _contains_any(normalized, READING_KEYWORDS) or _looks_like_bibliography(normalized):
        return "lectura"
    if _is_actionable_activity_text(normalized):
        return "actividad"
    return "otro"


def _extract_title(line: str, matched_date_text: str, category: str) -> str:
    title = re.sub(re.escape(matched_date_text), "", line, flags=re.IGNORECASE).strip(" :.-")
    normalized = _normalize_match_text(title)
    if category == "lectura":
        return _summarize_reading_section(normalized, normalized)
    if category == "actividad":
        return _summarize_activity_segment(normalized, normalized)
    return ""


def _extract_subject(text: str, source_name: str) -> str:
    lines = _normalize_lines(text)
    header_candidates = lines[:20]

    patterns = [
        r"\b(?:materia|asignatura|curso|catedra|modulo)\s*:\s*(.+)",
        r"\b(?:plan de curso|programa de curso|syllabus)\s+de\s+(.+)",
        r"\bnombre del curso / seminario\s*\|\s*(.+)",
    ]

    for line in header_candidates:
        for pattern in patterns:
            match = re.search(pattern, _normalize_match_text(line), flags=re.IGNORECASE)
            if match:
                subject = _clean_subject(match.group(1))
                if subject:
                    return subject

    stem = Path(source_name).stem if source_name else ""
    cleaned_name = _clean_subject(stem.replace("_", " ").replace("-", " "))
    return cleaned_name or "Materia no identificada"


def _extract_subject_from_document(document: Document, source_name: str) -> str:
    for table in document.tables:
        for row in table.rows:
            cells = _unique_nonempty(_row_to_cells(row))
            joined = " | ".join(cells)
            normalized = _normalize_match_text(joined)
            match = re.search(r"nombre del curso / seminario\s*\|\s*(.+)", normalized, flags=re.IGNORECASE)
            if match:
                subject = _clean_subject(match.group(1))
                if subject:
                    return subject

    text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    return _extract_subject(text, source_name)


def _extract_year_from_document(document: Document, source_name: str, fallback_year: int) -> int:
    source_match = re.search(r"\b(20\d{2})[-_ ]?[12]?\b", Path(source_name).stem)
    if source_match:
        return int(source_match.group(1))

    for table in document.tables:
        for row in table.rows:
            joined = " | ".join(_unique_nonempty(_row_to_cells(row)))
            normalized = _normalize_match_text(joined)
            match = re.search(r"periodo academico\s*\|\s*(20\d{2})", normalized, flags=re.IGNORECASE)
            if match:
                return int(match.group(1))

    return fallback_year


def _unique_nonempty(values: list[str]) -> list[str]:
    compact: list[str] = []
    for value in values:
        if value and value not in compact:
            compact.append(value)
    return compact


def _clean_subject(value: str) -> str:
    cleaned = _normalize_match_text(value)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" :.-_")
    cleaned = re.sub(
        r"\b(grupo|semestre|periodo|nivel|modalidad)\b.*$",
        "",
        cleaned,
        flags=re.IGNORECASE,
    ).strip(" :.-_")
    return _truncate_display(cleaned.title(), limit=120)


def _clean_detail_text(value: str) -> str:
    cleaned = _normalize_match_text(value)
    cleaned = re.sub(r"https?://\S+", "", cleaned)
    cleaned = re.sub(r"\d+\s*%", "", cleaned)
    cleaned = re.sub(r"\bpp?\.\s*\d+[^\s,;]*", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" :.,;-_")
    cleaned = re.sub(r"^(?:de|del|la|el|los|las|sobre|para|por|con|en)\s+", "", cleaned)
    return cleaned.strip(" :.,;-_")


def _extract_theme_hint(text: str) -> str:
    normalized = _normalize_match_text(text)
    parts = re.split(r"[.;|]", normalized)
    for part in parts:
        cleaned = _clean_detail_text(_strip_section_prefix(part))
        if not cleaned:
            continue
        if _contains_any(cleaned, NOISE_KEYWORDS):
            continue
        if _contains_any(cleaned, STRONG_ACTIVITY_KEYWORDS):
            continue
        if len(cleaned.split()) < 2:
            continue
        return _truncate_display(cleaned, limit=90)
    return ""


def _extract_first_reading_phrase(section: str) -> str:
    parts = re.split(r"[.;|]", section)
    for part in parts:
        cleaned = _clean_detail_text(part)
        if not cleaned:
            continue
        if _contains_any(cleaned, NOISE_KEYWORDS):
            continue
        if len(cleaned.split()) < 3:
            continue
        if cleaned.count(",") >= 1 and ":" not in cleaned and len(_extract_significant_words(cleaned)) <= 2:
            continue
        return cleaned
    return ""


def _truncate_display(value: str, limit: int = 100) -> str:
    compact = re.sub(r"\s+", " ", value).strip()
    if len(compact) <= limit:
        return compact
    truncated = compact[: limit - 3].rsplit(" ", 1)[0].strip()
    return f"{truncated}..."


def _strip_section_prefix(text: str) -> str:
    return re.sub(
        r"^(?:desarrollo(?:\s+y\s+neuropsicologia)?|neuropsicologia(?:\s+y\s+desarrollo)?|psicopatologia|psicologia clinica)\s*:\s*",
        "",
        text,
        flags=re.IGNORECASE,
    ).strip()


def _contains_any(normalized_text: str, keywords: tuple[str, ...]) -> bool:
    return any(_keyword_in_text(normalized_text, keyword) for keyword in keywords)


def _keyword_in_text(normalized_text: str, keyword: str) -> bool:
    pattern = r"\b" + re.escape(keyword).replace(r"\ ", r"\s+") + r"\b"
    return re.search(pattern, normalized_text) is not None


def _looks_like_bibliography(text: str) -> bool:
    markers = (
        "doi",
        "revista",
        "editorial",
        "capitulo",
        "capitulos",
        "http://",
        "https://",
        "arlington",
        "ciudad de mexico",
    )
    return any(marker in text for marker in markers)


def _dedupe_strings(values: list[str]) -> list[str]:
    seen: set[str] = set()
    unique: list[str] = []
    for value in values:
        compact = value.strip()
        if not compact:
            continue
        key = compact.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(compact)
    return unique


def _normalize_match_text(value: str) -> str:
    folded = unicodedata.normalize("NFKD", value)
    folded = "".join(char for char in folded if not unicodedata.combining(char))
    return folded.lower()
